from openai import OpenAI
from dotenv import load_dotenv
import os
import json
import math
import pandas as pd
pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', 100)
from difflib import SequenceMatcher
from docx import Document
from image_details_extractor import generate_product_description
from OLD.analytics_matcher import match_headline_to_keyword
import re
from typing import List, Iterable

# Load environment variables from .env file
load_dotenv()

# Initialize the OpenAI client using your endpoint and token
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
)

# Sheets: ['Model Training', '1 New Romance Copy Generation', '2 New Products Part of MegaPDP\u200b', '3 Products for SEO Enrichment', 'Full Catalog Short Descriptions']
# Load Coach Rules
coach_doc = Document("Documents/Coach Rules_new.docx")
coach_rules = "\n".join([para.text for para in coach_doc.paragraphs])

# Load Kate Spade Rules
spade_doc = Document("Documents/Kate Spade Rules.docx")
spade_rules = "\n".join([para.text for para in spade_doc.paragraphs])

# Data File paths
file1 = 'Documents/POC Product Selection- Coach Outlet.xlsx'
file2 = 'Documents/POC Product Selection- Kate Spade.xlsx'

coach_sheets = pd.read_excel(file1, sheet_name=None)
spade_sheets = pd.read_excel(file2, sheet_name=None)

def similarity(a: str, b: str) -> float:
    """Return a float [0.0–1.0] for how similar two strings are."""
    return SequenceMatcher(None, a, b).ratio()

def get_tagline(product_attributes, company):
    raw_mega_value = product_attributes.get("Mega PDP Group Value", "")
    mega_value = str(raw_mega_value).lower() if pd.notna(raw_mega_value) else ""

    rules= coach_rules
    mega_unique_values = (
        set(val.lower() for val in coach_sheets["Model Training"]["Mega PDP Group Value"].dropna().unique())
        .union(
            val.lower() for val in coach_sheets["Full Catalog Short Descriptions"]["Mega PDP Group Value"].dropna().unique()
        )
    )

    prod_old_description = []
    matched_mega_values = []
    match_type = ""

    if mega_value in mega_unique_values:
        print("Exact match")
        match_type = "Exact"
        for sheet_name in ("Model Training", "Full Catalog Short Descriptions"):
            df = coach_sheets[sheet_name]
            matches = df[df["Mega PDP Group Value"].str.lower() == mega_value]
            if not matches.empty:
                prod_old_description.extend(matches["Short Description - en"].tolist())
                matched_mega_values.extend(matches["Mega PDP Group Value"].tolist())
    
    else:
        similarity_scores = [(val, similarity(mega_value, val)) for val in mega_unique_values]
        similarity_scores.sort(key=lambda x: x[1], reverse=True)
        
        if similarity_scores:
            def fetch_matches(match_vals):
                for match_val in match_vals:
                    for sheet_name in ("Model Training", "Full Catalog Short Descriptions"):
                        df = coach_sheets[sheet_name]
                        matches = df[df["Mega PDP Group Value"].str.lower() == match_val]
                        if not matches.empty:
                            prod_old_description.extend(matches["Short Description - en"].tolist())
                            matched_mega_values.extend(matches["Mega PDP Group Value"].tolist())
                            return True  # Stop at first valid match
                return False

            top_match_val, top_match_score = similarity_scores[0]

            if top_match_score >= 0.7:
                print(f"Found similar match: {top_match_val} (Score: {top_match_score})")
                match_type = "Similar threshold greater than 70%"
                fetch_matches([top_match_val])
            else:
                print("No strong match found, returning top 10 matches")
                match_type = "Less than 70%"
                top_10_match_vals = [val for val, _ in similarity_scores[:10]]
                for val, score in similarity_scores[:10]:
                    print(f"Match: {val} (Score: {score})")
                fetch_matches(top_10_match_vals)
    blacklisted_keywords = [
        "inspired by", "chic", "exudes sophistication", "gen-z customer", "gen-z", 
        "aesthetic", "affordable", "ageless", "body", "chic", "coachie", "couture", 
        "craftsman", "customer", "cute", "dainty", "daintier", "darling", "deal", 
        "delightful", "designer", "discount", "disruptive", "don", "donning", 
        "easy win", "elegant", "elegance", "embellished", "enchanting", "engineered", 
        "eternal", "expressive luxury", "fabulous", "fabulousness", "fashion", 
        "fashion lover", "fashionista", "fave", "footwear", "gang", "gender-neutral", 
        "handbag", "hot", "it bag", "it girl", "it’s giving", "jet-set", "lovely", 
        "multifunctional", "must have", "new you", "obsess", "obsessed", "mindful", 
        "green", "conscious", "eco-conscious", "pioneering", "pleasing", "pretty", 
        "purse", "quiet luxury", "sale", "sassy", "savage", "sensations", "sleek", 
        "splendid", "sueded", "sustainable", "szn", "tender", "treasures", 
        "trendsetter", "turn heads", "unearth", "unveil", "unveiling", "uptown style", 
        "downtown style", "urban", "vibes", "but make it fashion", "meet", 
        "experience", "introducing", "just", "literally", "figuratively", 
        "audacious", "pvc", "PVC", "mundane", "nitty-gritties", "beauty scores", 
        "best", "boast", "booster", "statement", "promise", "declaration", "go-to", 
        "taste", "impeccable", "pretty face", "testament", "touches", "must-have", 
        "impraczcal", "unassuming", "overlook", "unusual", "friend", "flair", 
        "fierce", "efforzless", "glamour", "outing", "fashionable", "stylish", 
        "more than a pretty face", "your new best friend", "accessories collection", 
        "boasts", "modern fashion", "this is Coach Outlet's promise to you", 
        "declaration of style", "testament to your impeccable taste", 
        "finishing touches from Coach Outlet", "fashion adventures", "flair", 
        "audacious modern style", "this beauty scores high", "meziculously", 
        "captivating", "aesthetics", "simplicity and class", "dash of the unusual", 
        "crafted to fulfill", "unassuming elegance", "it's impractical to overlook", 
        "sexy", "let's talk about", "inspiration can come", "fall in love", 
        "inspiration can strike", "picture this", "picture themselves", "imagine", 
        "bio-attributed", "bio-based", "biodegradable", "bio-finished", 
        "carbon neutral", "certified b corp", "chemical recycling", "circular", 
        "closed loop", "compostable", "fair trade", "FSC", "forest stewardship", 
        "council", "mechanical recycling", "natural", "PEFC", "recyclable", 
        "upcycled", "SFI", "responsible", "synthetic", "traceable", "transparent", 
        "vegan", "zero waste"
    ]

    prod_old_description = list(set(prod_old_description))

    def remove_blacklisted_keywords(paragraphs: Iterable[str], blacklisted_keywords: Iterable[str]) -> List[str]:
        pattern = r'\b(' + '|'.join(blacklisted_keywords) + r')\b'
        regex = re.compile(pattern, re.IGNORECASE)

        seen = set()
        cleaned_paragraphs = []
        for para in paragraphs:
            # Skip invalid entries
            if not isinstance(para, str) or not para.strip():
                continue

            # Single regex substitution for all keywords
            cleaned = regex.sub('', para)
            # Normalize whitespace
            cleaned = ' '.join(cleaned.split())
            
            # Skip duplicates (case-insensitive)
            cleaned_lower = cleaned.lower()
            if cleaned_lower not in seen:
                seen.add(cleaned_lower)
                cleaned_paragraphs.append(cleaned)

        return cleaned_paragraphs
    
    prod_old_description = remove_blacklisted_keywords(prod_old_description, blacklisted_keywords)

    prompt = [
    "Instructions:",
    "1. You are given a set of rules. Follow them exactly to generate a new product description.",
    "** The product description should appeal to Gen Z customers and sound natural and effortless.**"
    "2. Do NOT repeat any word in the product description.",
    "3. Do NOT include the phrase “what fits inside.”",
    "4. Do NOT include any city-specific references.",
    "5. Do NOT mention pairing dress or attire.",
    "6. Do NOT use general phrases—be product-specific.",
    "7. Maintain a natural, authentic tone.",
    "8. Historical Artifacts should be mentioned from the sample description only.",
    "9. Strictly avoid all blacklisted words (severe penalty for violations).",
    "10. Use Natural and Effortless tone. Do NOT use generic terms",
    "11. Content in product description about the product should be in third party, and address customer as 'you'. Remember 'you' must be used maximum of once per product description",
    "**You must follow Instructions and rules at any cost, else you will be heavily penalized.**",
    "####",
    "Rules:",
    f"{rules}",
    "####",
    ]

    if prod_old_description != "":
        prompt.append("**Use the sample description below as a guide to frame the product description. Ensure the new product description mirrors the tone and style of the sample. Do not use blacklisted words. Only include historical artifacts if they are mentioned in the sample.**")
        prompt.append(f"{prod_old_description}")

    prompt += [
        "####",
        "SEO Guidance:",
        "Generate an SEO keyword list.",
        "Keywords should be specific to the product and not generic.",
        "Keyword Hierarchy (incorporate these into your product description where natural):",
        "- **Primary Keywords** - Describing product type (example - Flap shoulder bag, Colorblocked bag, Convertible bag)",
        "- **Secondary Keywords** - Describing characteristics (example - Pebbled leather, Colorblocked leather, Classic flap silhouette, Adjustable crossbody strap, Convertible design)",
        "- **Tertiary Keywords** - Describing function - (example - Optional crossbody strap, Everyday bag, Versatile handbag)",
    ]

    # if product_description_image != {}:
    #     prompt.extend(["####"," Below is the visual description of the image. Do take into account while framing the product description.",
    #         f"{product_description_image}"])
    
    prompt.append("####")
    prompt.append("Below are the attributes for the product:")

    blacklist_pattern = re.compile(r'\b(?:' + '|'.join(re.escape(word) for word in blacklisted_keywords) + r')\b', re.IGNORECASE)

    for key, value in product_attributes.items():
        if key in ["What Fits Inside - en", "Iteration", "Tech Fit - en", "Primary Digital Asset URL", "Non-Primary Digital Asset URL"]:
            continue
        if value == "" or (isinstance(value, float) and math.isnan(value)):
            continue

        if not isinstance(value, str):
            value = json.dumps(value, ensure_ascii=False)

        cleaned_value = blacklist_pattern.sub("", value)

        cleaned_value = re.sub(r'\s{2,}', ' ', cleaned_value).strip()

        pretty_value = json.dumps(cleaned_value, indent=2, ensure_ascii=False)
        indented_value = "\n".join([f"  {line}" for line in pretty_value.splitlines()])
        prompt.append(f"- {key}:\n{indented_value}")


    prompt.extend([
        "####",
        f"Before generating, STRICTLY ENSURE none of the following blacklisted words appear in ANY output.",
        "\n",
        f"**Blacklisted Words (JSON array)**: **{json.dumps(blacklisted_keywords, ensure_ascii=False)}**",
        "####",
        "Do NOT use any blacklisted keywords in the product description.",
        "Always generate product description in 4 sentences described.",
    ])

    prompt.extend([
    "Format your response exactly like this (so it’s easy to parse):\n",
    "```json",
    "{",
    '  "product_description": "...",',
    '  "SEO Keyword 1": ["...", "...", "..."],',
    '  "SEO Keyword 2": ["...", "...", "..."],',
    '  "SEO Keyword 3": ["...", "...", "..."]',
    "}",
    "```",
    ])

    full_prompt = "\n".join(prompt)

    system_prompts = [
    f"You are a world-class luxury fashion editor for {company}. Do NOT add Blacklisted Words in the product description."
    "Instructions:",
    "1. You are given a set of rules. Follow them exactly to generate a new product description.",
    "2. Do NOT repeat any word in the product description.",
    "3. Do NOT include the phrase “what fits inside.”",
    "4. Do NOT include any city-specific references.",
    "5. Do NOT mention pairing dress or attire.",
    "6. Do NOT use general phrases—be product-specific.",
    "7. Maintain a natural, authentic tone.",
    "8. Historical Artifacts should be mentioned from the sample description only.",
    "9. Strictly avoid all blacklisted words (severe penalty for violations).",
    "10. Use Natural and Effortless tone. Do NOT use generic terms",
    "11. Content in product description about the product should be in third party, and address customer as 'you'. Remember 'you' must be used maximum of once per product description",
    "**You must follow Instructions and rules at any cost, else you will be heavily penalized.**",
    "####",
    f"Blacklisted Keywords: {blacklisted_keywords}"
    ]

    system_prompts = "\n".join(system_prompts)

    # print(full_prompt)

    try:

        response = client.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": system_prompts},
                    {"role": "user", "content": full_prompt}
                ],
                temperature=0.3,
                response_format={"type": "json_object"}
            )
        
        res = json.loads(response.choices[0].message.content.strip())

        tagline = ''.join([
            res["product_description"],
        ])

        def count_sentences(text):
            # This pattern looks for '.', '!', or '?' followed by space or end of string
            sentences = re.split(r'[.!?](?:\s|$)', text.strip())
            # Remove empty strings
            sentences = [s for s in sentences if s.strip()]
            return len(sentences)

        # Check each key
        for key in ["product_description"]:
            text = res.get(key, "")
            count = count_sentences(text)
            print(f"{key}: {'✅ Has 4 sentences' if count == 4 else f'❌ Has {count} sentence(s)'}")
        
        # Create single regex pattern with all keywords
        pattern = r'\b(?:' + '|'.join(map(re.escape, blacklisted_keywords)) + r')\b'
        
        # Find all matches in one go
        found_keywords = re.findall(pattern, tagline, re.IGNORECASE)
        
        if found_keywords:
            print(f"🚨 Blacklisted keywords present: {', '.join(set(found_keywords))}")
        else:
            print("No blacklisted keywords found.")
    

        res["Old Description"] = prod_old_description
        res["Matched OLD Mega PDP Value"] = matched_mega_values
        res["Prompt"] = full_prompt
        res["Match_Type"] = match_type
        res['Blacklisted Keywords'] = found_keywords
        print("#####")
        print("product_description")
        print(res["product_description"])
        print("#####")
        return res
    
    except Exception as e:
        print("Error occurued")
        return {"product_description":{}}

def process_usecase(usecase_df, brand):
    data = usecase_df.to_dict(orient='records')
    output_data = []
    
    for item in data[:14]:
        print(f"Processing {item['Item#']}")
        # image = item.get("Primary Digital Asset URL", "")
        # image2 = item.get("Primary Digital Asset URL", "")  # Note: This might need adjustment if a secondary image column exists
        # raw_urls = f"{image}`{image2}".replace("`", ",").split(",")
        # images = list(filter(None, map(str.strip, raw_urls)))
        
        # if images:
        #     product_description_image = generate_product_description(images)
        # else:
        #     product_description_image = {}
        
        luxury_tagline = get_tagline(item, brand)
        
        if isinstance(luxury_tagline, dict):
            for k, v in luxury_tagline.items():
                item[k] = v
        else:
            item["product_description"] = luxury_tagline
        
        output_data.append(item)
    
    return pd.DataFrame(output_data)





